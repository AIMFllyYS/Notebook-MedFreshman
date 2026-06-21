"""
DOCX → Markdown 降级解析脚本（MinerU 不可用时使用）

用法：
    python scripts/fallback-docx.py <input.docx> <output.md>

依赖安装（二选一）：
    方案 A: pip install python-docx
    方案 B: 系统安装 pandoc（精度更高）

优先使用 pandoc（若可用），否则回退到 python-docx 纯文本提取。
"""

import sys
import os
import subprocess
import shutil


def docx_via_pandoc(docx_path: str, output_path: str) -> bool:
    """尝试使用 pandoc 转换，返回是否成功"""
    if not shutil.which("pandoc"):
        return False
    try:
        subprocess.run(
            ["pandoc", docx_path, "-o", output_path, "--wrap=none"],
            check=True,
            capture_output=True,
        )
        print(f"✓ {os.path.basename(docx_path)} → {output_path} (via pandoc)")
        return True
    except subprocess.CalledProcessError as e:
        print(f"  pandoc failed: {e.stderr.decode()}")
        return False


def docx_via_python_docx(docx_path: str, output_path: str) -> None:
    """使用 python-docx 提取纯文本"""
    try:
        from docx import Document
    except ImportError:
        print("ERROR: Neither pandoc nor python-docx available.")
        print("  Install one of:")
        print("    pip install python-docx")
        print("    Or install pandoc: https://pandoc.org/installing.html")
        sys.exit(1)

    doc = Document(docx_path)
    md_lines: list[str] = []
    md_lines.append(f"# {os.path.basename(docx_path)}\n")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_lines.append("")
            continue

        style_name = (para.style.name or "").lower()
        if "heading 1" in style_name:
            md_lines.append(f"## {text}\n")
        elif "heading 2" in style_name:
            md_lines.append(f"### {text}\n")
        elif "heading 3" in style_name:
            md_lines.append(f"#### {text}\n")
        elif "list" in style_name:
            md_lines.append(f"- {text}")
        else:
            md_lines.append(text)

    for table in doc.tables:
        md_lines.append("")
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            md_lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                md_lines.append("| " + " | ".join("---" for _ in cells) + " |")
        md_lines.append("")

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))
    print(f"✓ {os.path.basename(docx_path)} → {output_path} (via python-docx)")


def main():
    if len(sys.argv) != 3:
        print("Usage: python scripts/fallback-docx.py <input.docx> <output.md>")
        sys.exit(1)

    docx_path = sys.argv[1]
    output_path = sys.argv[2]

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

    if not docx_via_pandoc(docx_path, output_path):
        docx_via_python_docx(docx_path, output_path)


if __name__ == "__main__":
    main()
